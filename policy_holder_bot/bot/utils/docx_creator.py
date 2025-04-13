from docx import Document
from typing import Dict
from datetime import datetime

from bot.config.diagnosis_config import diagnosis_config

def create_docx(policy_number: str, diagnosis_code: str, diagnosis_date: str, crypto_wallet: str) -> str:
    doc = Document()
    
    # Add header
    doc.add_heading('Insurance Claim Form', 0)
    
    # Add policy information
    doc.add_paragraph(f'Policy Number: {policy_number}')
    doc.add_paragraph(f'Diagnosis Code: {diagnosis_code}')
    doc.add_paragraph(f'Diagnosis Date: {diagnosis_date}')
    doc.add_paragraph(f'Crypto Wallet: {crypto_wallet}')
    
    # Add diagnosis description
    diagnosis_description = diagnosis_config.get(diagnosis_code, 'Unknown diagnosis code')
    doc.add_paragraph(f'Diagnosis Description: {diagnosis_description}')
    
    # Save document
    filename = f'claim_{policy_number}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    doc.save(filename)
    
    return filename